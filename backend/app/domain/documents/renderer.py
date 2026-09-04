from __future__ import annotations

import io
from dataclasses import dataclass
from enum import StrEnum

from app.core.config import get_settings
from app.domain.resume.extractor import ResumeExtraction

_SECTION_ORDER = ("summary", "skills", "experience", "projects", "education", "certifications")

_HTML_STYLE = (
    "body{font-family:Georgia,'Times New Roman',serif;max-width:760px;"
    "margin:2rem auto;color:#1a1a1a;line-height:1.5}"
    "h1{margin-bottom:0.2rem}"
    "h2{border-bottom:1px solid #ccc;padding-bottom:0.2rem;margin-top:1.5rem}"
    "h3{margin-bottom:0.2rem}"
    "ul{margin-top:0.3rem}"
)


class RenderFormat(StrEnum):
    MD = "md"
    HTML = "html"
    PDF = "pdf"
    DOCX = "docx"


@dataclass(frozen=True)
class RenderedDoc:
    fmt: RenderFormat
    media_type: str
    data: bytes


class RenderUnavailable(RuntimeError):
    pass


def _date_range(start: str | None, end: str | None, *, is_current: bool = False) -> str:
    if is_current:
        end = "Present"
    if start and end:
        return f"{start} - {end}"
    if start:
        return start
    if end:
        return end
    return ""


class DocumentRenderer:
    def to_markdown(self, r: ResumeExtraction) -> str:
        lines: list[str] = [f"# {r.full_name or 'Unnamed Candidate'}"]

        contact_parts = [
            p
            for p in (r.email, r.location, r.github_url, r.linkedin_url, r.portfolio_url)
            if p
        ]
        if contact_parts:
            lines.append(" | ".join(contact_parts))

        for section in _SECTION_ORDER:
            if section == "summary":
                if r.summary:
                    lines.append("")
                    lines.append("## Summary")
                    lines.append(r.summary)
            elif section == "skills":
                if r.skills:
                    lines.append("")
                    lines.append("## Skills")
                    lines.append(", ".join(r.skills))
            elif section == "experience":
                if r.experiences:
                    lines.append("")
                    lines.append("## Experience")
                    for exp in r.experiences:
                        lines.append("")
                        lines.append(f"### {exp.title}, {exp.company}")
                        date_range = _date_range(
                            exp.start_date, exp.end_date, is_current=exp.is_current
                        )
                        if date_range:
                            lines.append(date_range)
                        if exp.description:
                            lines.append(exp.description)
                        for highlight in exp.highlights:
                            lines.append(f"- {highlight}")
            elif section == "projects":
                if r.projects:
                    lines.append("")
                    lines.append("## Projects")
                    for proj in r.projects:
                        lines.append("")
                        lines.append(f"### {proj.name}")
                        date_range = _date_range(proj.start_date, proj.end_date)
                        if date_range:
                            lines.append(date_range)
                        if proj.description:
                            lines.append(proj.description)
                        for highlight in proj.highlights:
                            lines.append(f"- {highlight}")
            elif section == "education":
                if r.education:
                    lines.append("")
                    lines.append("## Education")
                    for edu in r.education:
                        lines.append("")
                        degree_field = ", ".join(p for p in (edu.degree, edu.field) if p)
                        heading = f"### {edu.institution}"
                        if degree_field:
                            heading = f"### {degree_field}, {edu.institution}"
                        lines.append(heading)
                        date_range = _date_range(edu.start_date, edu.end_date)
                        if date_range:
                            lines.append(date_range)
                        if edu.grade:
                            lines.append(edu.grade)
            elif section == "certifications":
                if r.certifications:
                    lines.append("")
                    lines.append("## Certifications")
                    for cert in r.certifications:
                        lines.append("")
                        heading = f"### {cert.name}"
                        if cert.issuer:
                            heading = f"### {cert.name}, {cert.issuer}"
                        lines.append(heading)
                        date_range = _date_range(cert.issued_date, cert.expires_date)
                        if date_range:
                            lines.append(date_range)

        return "\n".join(lines) + "\n"

    def to_html(self, r: ResumeExtraction) -> str:
        try:
            from markdown_it import MarkdownIt
        except ImportError as exc:
            raise RenderUnavailable("markdown_it is not available") from exc

        body = MarkdownIt().render(self.to_markdown(r))
        return f"<html><head><style>{_HTML_STYLE}</style></head><body>{body}</body></html>"

    def to_pdf(self, r: ResumeExtraction) -> bytes:
        if not get_settings().doc_render_enabled:
            raise RenderUnavailable("document rendering is disabled")

        try:
            from xhtml2pdf import pisa
        except ImportError as exc:
            raise RenderUnavailable("xhtml2pdf is not available") from exc

        buf = io.BytesIO()
        result = pisa.CreatePDF(self.to_html(r), dest=buf)
        if result.err:
            raise RenderUnavailable("xhtml2pdf failed to render the document")
        return buf.getvalue()

    def to_docx(self, r: ResumeExtraction) -> bytes:
        if not get_settings().doc_render_enabled:
            raise RenderUnavailable("document rendering is disabled")

        try:
            from docx import Document
        except ImportError as exc:
            raise RenderUnavailable("python-docx is not available") from exc

        doc = Document()
        doc.add_heading(r.full_name or "Unnamed Candidate", level=1)

        contact_parts = [
            p
            for p in (r.email, r.location, r.github_url, r.linkedin_url, r.portfolio_url)
            if p
        ]
        if contact_parts:
            doc.add_paragraph(" | ".join(contact_parts))

        if r.summary:
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(r.summary)

        if r.skills:
            doc.add_heading("Skills", level=2)
            doc.add_paragraph(", ".join(r.skills))

        if r.experiences:
            doc.add_heading("Experience", level=2)
            for exp in r.experiences:
                doc.add_heading(f"{exp.title}, {exp.company}", level=3)
                date_range = _date_range(exp.start_date, exp.end_date, is_current=exp.is_current)
                if date_range:
                    doc.add_paragraph(date_range)
                if exp.description:
                    doc.add_paragraph(exp.description)
                for highlight in exp.highlights:
                    doc.add_paragraph(highlight, style="List Bullet")

        if r.projects:
            doc.add_heading("Projects", level=2)
            for proj in r.projects:
                doc.add_heading(proj.name, level=3)
                date_range = _date_range(proj.start_date, proj.end_date)
                if date_range:
                    doc.add_paragraph(date_range)
                if proj.description:
                    doc.add_paragraph(proj.description)
                for highlight in proj.highlights:
                    doc.add_paragraph(highlight, style="List Bullet")

        if r.education:
            doc.add_heading("Education", level=2)
            for edu in r.education:
                degree_field = ", ".join(p for p in (edu.degree, edu.field) if p)
                heading = f"{degree_field}, {edu.institution}" if degree_field else edu.institution
                doc.add_heading(heading, level=3)
                date_range = _date_range(edu.start_date, edu.end_date)
                if date_range:
                    doc.add_paragraph(date_range)
                if edu.grade:
                    doc.add_paragraph(edu.grade)

        if r.certifications:
            doc.add_heading("Certifications", level=2)
            for cert in r.certifications:
                heading = f"{cert.name}, {cert.issuer}" if cert.issuer else cert.name
                doc.add_heading(heading, level=3)
                date_range = _date_range(cert.issued_date, cert.expires_date)
                if date_range:
                    doc.add_paragraph(date_range)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def render(self, r: ResumeExtraction, fmt: RenderFormat) -> RenderedDoc:
        if fmt is RenderFormat.MD:
            data = self.to_markdown(r).encode()
            return RenderedDoc(fmt=fmt, media_type="text/markdown", data=data)
        if fmt is RenderFormat.HTML:
            return RenderedDoc(fmt=fmt, media_type="text/html", data=self.to_html(r).encode())
        if fmt is RenderFormat.PDF:
            return RenderedDoc(fmt=fmt, media_type="application/pdf", data=self.to_pdf(r))
        if fmt is RenderFormat.DOCX:
            media_type = (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            return RenderedDoc(fmt=fmt, media_type=media_type, data=self.to_docx(r))
        raise ValueError(f"unsupported render format: {fmt!r}")
